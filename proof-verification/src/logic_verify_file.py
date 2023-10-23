import itertools
import pathlib
from abc import abstractmethod
from itertools import zip_longest
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from docx import *
import logic_rules as proof
from tkinter.filedialog import askopenfilename
from shutil import which
import shlex
import configparser
import os
import re

langmap = configparser.ConfigParser()


def always_valid(*args, **kwargs):
    return True


def always_invalid(*args, **kwargs):
    return False


def ignore(*args, **kwargs):
    return None


reason_map = {
    "∨O": proof.or_out,
    "∨I": proof.or_in,
    "~∨O": proof.not_or_out,
    "&O": proof.and_out,
    "&I": proof.and_in,
    "~&O": proof.not_and_out,
    "→O": proof.if_out,
    "~→O": proof.not_if_out,
    "↔O": proof.iff_out,
    "↔I": proof.iff_in,
    "~↔O": proof.not_iff_out,
    "DN": proof.double_negation,
    "REP": proof.rep,
    "⨳I": proof.x_in,
    "⨳O": always_valid,
    "∀O": proof.all_out,
    "~∀O": proof.not_all_out,
    "∃O": proof.exist_out,
    "∃I": proof.exist_in,
    "~∃O": proof.not_exist_out,
    "PR": always_valid,
    "AS": ignore,
}

show_map = {
    "DD": (0, always_valid),
    "~D": (2, proof.show_td),
    "ID": (2, proof.show_id),
    "CD": (2, proof.show_cd),
    "∀D": (1, proof.show_ud),
    "UD": (1, proof.show_ud)
}


def extract_reason(line):
    line = clean_language(line)
    rule = re.findall(r"[^0-9\ \,\.]+", line)
    sources = re.findall(r"[0-9]+", line)
    verification_method = reason_map.get(rule[0].upper(), always_invalid)
    return [int(i) for i in sources], verification_method


def extract_show_follow(tp):
    return show_map[tp.upper()]


def clean_language(line):
    if not line:
        return line
    if "symbols" not in langmap:
        return line
    for operand, symbol in langmap["symbols"].items():
        line = line.replace(symbol, proof.default_language[operand])
    return line


def clean_line(line):
    if not line:
        return line
    line = clean_language(line)
    return line.replace("[", "(").replace("]", ")").replace(" ", "").rstrip()


def find_start(column):
    for index, cell in enumerate(column):
        if cell.value is not None:
            return index + 1


class Proof:
    def __init__(self):
        self._index = 1
        self.show_regex = re.compile("show:?", re.IGNORECASE)

    def handle_proof(self):
        while self.get_line(self._index)[0]:
            row = self.get_line(self._index)
            if not isinstance(row[0], int) or row[1] is None:
                self._index += 1
                continue
            try:
                if "SHOW" in row[1].upper():
                    self.verify_showline()
                else:
                    self.verify_proofline()
            except:
                self.mark_error()
            self._index += 1

    def verify_showline(self):
        _, line, method = self.get_line(self._index)
        line = self.show_regex.sub("", line)
        if method is None or method == "":
            self.mark_unknown()
            return
        show = proof.shunting_yard_statement(line)
        following_size, rule = extract_show_follow(method)
        following = [self.get_line(self._index + l + 1)[1] for l in range(following_size)]
        following = [self.show_regex.sub("", l) for l in following]
        following = [proof.shunting_yard_statement(l) for l in following]

        if len(following) <= 1:
            valid = rule(show, *following)
        else:
            valid = rule(show, following)

        if valid is None:
            return
        elif not valid:
            for line in range(following_size + 1):
                self.mark_failure(self._index + line)
        else:
            for line in range(following_size + 1):
                self.mark_success(self._index + line)

    def verify_proofline(self):
        _, line, reason = self.get_line(self._index)
        if reason is None or reason == "":
            self.mark_unknown()
            return
        relies, rule = extract_reason(reason)
        relies = list(map(proof.shunting_yard_statement,
                          map(lambda row: self.show_regex.sub("", row),  # this could use a closed showline
                              map(lambda row: self.get_line(row)[1], relies))))
        if len(relies) <= 1:
            valid = rule(*relies, proof.shunting_yard_statement(line))
        else:
            valid = rule(relies, proof.shunting_yard_statement(line))
        if valid is None:
            return
        elif not valid:
            self.mark_failure()
        else:
            self.mark_success()

    def get_line(self, index):
        line = self._get_raw_line(index)
        line[1] = clean_line(line[1])
        line[2] = clean_language(line[2])
        if isinstance(line[0], str) and line[0].isdigit():
            line[0] = int(line[0])
        return line

    @abstractmethod
    def _get_raw_line(self, index):
        pass

    @abstractmethod
    def mark_error(self, index=None):
        pass

    @abstractmethod
    def mark_unknown(self, index=None):
        pass

    @abstractmethod
    def mark_failure(self, index=None):
        pass

    @abstractmethod
    def mark_success(self, index=None):
        pass


class ExcelProof(Proof):
    ERROR = PatternFill(start_color="2d0c45", end_color="2d0c45", fill_type="solid")
    RED = PatternFill(start_color="00FF0000", end_color="00FF0000", fill_type="solid")
    GREEN = PatternFill(start_color="0000FF00", end_color="0000FF00", fill_type="solid")
    YELLOW = PatternFill(start_color="FFFF0000", end_color="FFFF0000", fill_type="solid")

    def __init__(self, sheet, start=0):
        self.sheet = sheet
        self.start = start
        super().__init__()

    def mark_unknown(self, index=None):
        index = index or self._index
        self.sheet.cell(self.start + index, 3).fill = self.YELLOW

    def mark_failure(self, index=None):
        index = index or self._index
        self.sheet.cell(self.start + index, 3).fill = self.RED

    def mark_success(self, index=None):
        index = index or self._index
        self.sheet.cell(self.start + index, 3).fill = self.GREEN

    def mark_error(self, index=None):
        index = index or self._index
        self.sheet.cell(self.start + index, 3).fill = self.ERROR

    def _get_raw_line(self, index):
        return [self.sheet.cell(self.start + index, i + 1).value for i in range(3)]


class TextProof(Proof):
    def __init__(self, text, marking_format="{0} | {1}"):
        super().__init__()
        self.text = text
        self.line_regex = re.compile(r"(\d*)\.?[.,|\s]*([^0-9.,|]*)[,.|\s]+(.*)?")
        self.markings = [""] * len(text)
        self.marking_format = marking_format

    def mark_error(self, index=None):
        index = index or self._index
        self.markings[index - 1] = "!!!"

    def mark_unknown(self, index=None):
        index = index or self._index
        self.markings[index - 1] = "?"

    def mark_failure(self, index=None):
        index = index or self._index
        self.markings[index - 1] = "X"

    def mark_success(self, index=None):
        index = index or self._index
        self.markings[index - 1] = "V"

    def _get_raw_line(self, index):
        if index > len(self.text):
            return [None, None, None]
        line = self.line_regex.match(self.text[index - 1])
        if not line:
            return [None, None, None]
        return list(line.groups()[:3])

    def get_marked(self):
        return [self.marking_format.format(a, b) if b else a for a, b in zip(self.text, self.markings)]


class DocxProof(Proof):
    def __init__(self, table):
        super().__init__()
        self.table = table
        self.is_number = re.compile("\d+")
        self.start = 0 if self.is_number.match(table.cell(0, 0).text) else 1

    def mark_error(self, index=None):
        self._color_cell(((index or self._index) + self.start - 1, 2), "2d0c45")

    def mark_unknown(self, index=None):
        self._color_cell(((index or self._index) + self.start - 1, 2), "FFFF0000")

    def mark_failure(self, index=None):
        self._color_cell(((index or self._index) + self.start - 1, 2), "00FF0000")

    def mark_success(self, index=None):
        self._color_cell(((index or self._index) + self.start - 1, 2), "0000FF00")

    def _color_cell(self, index, color):
        shading_elm_1 = parse_xml(f'<w:shd {{}} w:fill="{color}"/>'.format(nsdecls('w')))
        try:
            self.table.cell(*index)._tc.get_or_add_tcPr().append(shading_elm_1)
        except IndexError:
            print(index)

    def _get_raw_line(self, index):
        if index + self.start > len(self.table.rows):
            return [None, None, None]
        line = [x.text for x in self.table.rows[index + self.start - 1].cells]
        line[0] = self.is_number.match(line[0]).group()
        return line


def find_indices(list_to_check, item_to_find):
    for idx, value in enumerate(list_to_check):
        if value == item_to_find:
            yield idx


def verify_excel(file):
    wb = load_workbook(file, data_only=True)
    sheet = wb.active
    indices = find_indices([cell.value for cell in sheet["A"]], 1)
    for start in indices:
        proof = ExcelProof(sheet, start)
        proof.handle_proof()
    name = file.split(".")[0]
    wb.save(f"{name}-checked.xlsx")


def verify_ods(f):
    libreoffice = which("libreoffice") or which("LibreOffice")
    if libreoffice is None:
        print("ERROR: LibreOffice is not installed but required for ods support")
        return False
    path = os.path.dirname(f)
    os.chdir(path)
    command = f"{libreoffice} --headless --invisible --convert-to xlsx {shlex.quote(f)}"
    print(command)
    os.system(command)
    converted = f.replace(".ods", ".xlsx")

    verify_excel(converted)

    checked = f.replace(".ods", "-checked.xlsx")
    command = (
        f"{libreoffice} --headless --invisible --convert-to ods {shlex.quote(checked)}"
    )
    print(command)
    os.system(command)
    os.remove(converted)
    os.remove(checked)
    return True


def verify_text(file):
    name = file.split(".")[0]
    checked = pathlib.Path(f"{name}-checked.txt")
    full_file = pathlib.Path(file).read_text().splitlines()
    start_line_regex = re.compile(r"^\s*1[\s.|]")  # the line starts with the number 1 and seperator
    indices = list(map(lambda x: x[0], filter(lambda x: start_line_regex.match(x[1]), enumerate(full_file))))
    ranges = zip_longest(indices, indices[1:])
    with checked.open("w") as checked:
        for start, end in ranges:
            proof = TextProof(full_file[start:end])
            proof.handle_proof()
            checked.writelines([p + "\n" for p in proof.get_marked()])
    return True


def verify_docx(file):
    doc = docx.Document(docx=file)
    if len(doc.tables) == 0:  # it is written in plain text instead of tables
        all_text = [paragraph.text.splitlines() for paragraph in doc.paragraphs]
        flattened = list(itertools.chain(*all_text))
        start_line_regex = re.compile(r"^\s*1[\s.|]")  # the line starts with the number 1 and seperator
        indices = list(map(lambda x: x[0], filter(lambda x: start_line_regex.match(x[1]), enumerate(flattened))))
        ranges = zip_longest(indices, indices[1:])
        with open(file.replace(".docx", "-checked.txt"), "w") as checked:
            for start, end in ranges:
                proof = TextProof(flattened[start:end])
                proof.handle_proof()
                checked.writelines([p + "\n" for p in proof.get_marked()])
        return True
    for proof in doc.tables:
        p = DocxProof(proof)
        p.handle_proof()
    print(file.replace(".docx", "-checked.docx"))
    doc.save(file.replace(".docx", "-checked.docx"))
    return True


def verify_file(f):
    if f.endswith(".xlsx"):
        verify_excel(f)
        return True
    if f.endswith(".ods"):
        return verify_ods(f)
    if f.endswith(".docx"):
        return verify_docx(f)
    if f.endswith(".txt"):
        return verify_text(f)
    return False


if True and __name__ == "__main__":
    configs = [os.getcwd(), os.path.join(os.path.expanduser("~"), ".config", "logic"),
               os.path.dirname(os.path.abspath(__file__)),
               "/etc", "/usr/local/etc", os.path.join(os.getenv("APPDATA", "", ), "logic")]
    configs = list(map(lambda x: os.path.join(x, "proof-verification.ini"), configs))
    configs = list(filter(os.path.exists, configs))
    langmap.read(configs)
    f = ""
    home_directory = os.path.expanduser("~")
    documents_directory = os.path.join(home_directory, "Documents")

    while not verify_file(f):
        f = askopenfilename(
            initialdir=documents_directory,
            filetypes=(("Spreadsheet files", "*.xlsx *.ods"),
                       ("Docs files", "*.docx"),
                       ("Raw Text", "*.txt"),
                       ("All Files", "*.*")),
        )
